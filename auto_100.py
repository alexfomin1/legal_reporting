import xlrd
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

print('!!! Формат названия отчета: Отчет 100 тыс. дд.мм.гггг.xlsx !!!')
print()
b_pr = input('Введите дату текущего отчета (дд.мм.гггг): ')
b_pr = 'Отчет 100 тыс. ' + b_pr + '.xlsx'

b_tod = input('Введите дату текущего отчета (дд.мм.гггг): ')
b_tod = 'Отчет 100 тыс. ' + b_tod + '.xlsx'

#введение таблицы
#b_tod = 'Отчет 100 тыс. 08.06.2022.xlsx'
book_tod = xlrd.open_workbook(b_tod)
sd = {}
for s in book_tod.sheets():
    sd[s.name] = s
sheet_tod = sd['Лист1']

#b_pr = 'Отчет 100 тыс. 01.06.2022.xlsx'
book_pr = xlrd.open_workbook(b_pr)
sd1 = {}
for s in book_pr.sheets():
    sd1[s.name] = s
sheet_pr = sd1['Лист1']

#табл -> число
def val_extr(val_t):
    a = str(val_t).split(':')
    a = float(a[1])
    a = float(f'{a:.2f}')
    return a

def dva(nach):
    nach = float(f'{nach:.2f}')
    return nach


#дата
date_tod = b_tod.split()
date_tod = date_tod[3].split('.')
del date_tod[-1]
date_tod = '.'.join(date_tod)

date_pr = b_pr.split()
date_pr = date_pr[3].split('.')
del date_pr[-1]
date_pr = '.'.join(date_pr)


#открытие документа
document = Document()
style = document.styles['Normal']
style.font.name = 'Times New Roman'
style.font.bold = True
style.font.size = Pt(12)
section = document.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Mm(10)
section.right_margin = Mm(5)
section.top_margin = Mm(2.5)
section.bottom_margin = Mm(5)
#paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#конструктор
#заголовок
head = document.add_heading('СПРАВКА')
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

head2 = document.add_heading('о состоянии работы по взысканию задолженности с потребителей, имеющих долг свыше 100 тыс. руб.', level=2)
head2.alignment = WD_ALIGN_PARAGRAPH.CENTER

#основной текст
amount_all = int(val_extr(sheet_tod.cell(40, 3)))
sum_all = val_extr(sheet_tod.cell(40, 6))
document.add_paragraph('По состоянию на {0} по данным установлено {1} ФЛС с задолженностью свыше 100 тыс. составила {2} млн. руб.'.format(date_tod, amount_all, sum_all))

para = document.add_paragraph('За период с 01.10.2020 по {} снижение составило:'.format(date_tod))

diff_fls = int(val_extr(sheet_tod.cell(40, 4)))
perc_diff_fls = val_extr(sheet_tod.cell(40, 5))
diff_din_fls = diff_fls - int(val_extr(sheet_pr.cell(40, 4)))
document.add_paragraph('-по количеству ФЛС {0} ({1}%); за неделю {2} ФЛС'.format(diff_fls, perc_diff_fls, diff_din_fls))
a = val_extr(sheet_tod.cell(40, 7))
b = dva(a / val_extr(sheet_tod.cell(40, 2)) * 100)
c = dva(a - val_extr(sheet_pr.cell(40, 7)))
d = dva(c / sum_all * 100)
document.add_paragraph('-по задолженности {0} млн. руб. ({1}%); за неделю {2} млн.руб. ({3}%)'.format(a, b, c, d))
a = val_extr(sheet_tod.cell(40, 9))
b = val_extr(sheet_tod.cell(40, 10))
document.add_paragraph('В настоящее время сумма задолженности, по которой ведется взыскание ГБУ Жилищник районов ЗАО, составляет {0} млн. руб. ({1}% от суммы задолженности).'.format(a, b))

run = document.add_paragraph().add_run('Наибольшее снижение задолженности в районах:')
run.font.color.rgb = RGBColor(0, 227, 49)

x, y = 27, 8
dictionary = {}
results = []
for i in range(13):
    dist = str(sheet_tod.cell(x, 0))
    dist = dist.split(':')
    dist = list(dist[1])
    dist = ''.join(dist[1:-1])  
    dictionary[val_extr(sheet_tod.cell(x, y))] = dist
    results.append(val_extr(sheet_tod.cell(x, y)))
    x += 1

results.sort()
n = 2
for i in range(3):
    a = dictionary[results[2 - i]]
    document.add_paragraph(a + ' - ' + str(-1 * results[2 - i]) + '%')

run = document.add_paragraph().add_run('Наименьшее снижение задолженности в районах:')
run.font.color.rgb = RGBColor(255, 23, 23)

results.sort(reverse=True)
n = 2
for i in range(3):
    a = dictionary[results[2 - i]]
    document.add_paragraph(a + ' - ' + str(-1 * results[2 - i]) + '%')
document.add_paragraph('')

#отсутствие динамики
x, y = 27, 8
no_din = []
fl = False
for i in range(13):
    if val_extr(sheet_tod.cell(x, y)) == val_extr(sheet_pr.cell(x, y)):
        fl = True
        no_din.append(x)
    x += 1
if fl == True:
    districts = []
    for x in no_din:
        dist = str(sheet_tod.cell(x, 0))
        dist = dist.split(':')
        dist = list(dist[1])
        dist = ''.join(dist[1:-1])
        districts.append(dist)        
if len(no_din) == 1:
    run = document.add_paragraph().add_run('Отсутствует динамика снижения задолженности за период с {0}-{1} в районе {2}.'.format(date_pr, date_tod, districts[0]))
    run.font.color.rgb = RGBColor(255, 23, 23)    
elif len(no_din) > 1:
    run = document.add_paragraph().add_run('Отсутствует динамика снижения задолженности за период с {0}-{1} в следующих районах:'.format(date_pr, date_tod))
    run.font.color.rgb = RGBColor(255, 23, 23)
    for x in districts:
        run = document.add_paragraph().add_run('- '+ x)
        run.font.color.rgb = RGBColor(255, 23, 23)

a = val_extr(sheet_tod.cell(40, 9))
document.add_paragraph('В разрезе проводимых мероприятий сумма задолженности, по которой ведется взыскание ({} млн. руб.), составляет:'.format(a))
a = val_extr(sheet_tod.cell(40, 11))
document.add_paragraph('- досудебные мероприятия – {} млн. руб. (ограничение водоотведения, проведены рейды по должникам, реструктуризация долга);'.format(a))
a = val_extr(sheet_tod.cell(40, 12))
document.add_paragraph('- судебно-исковое производство – {} млн. руб. (находятся на рассмотрении в суде);'.format(a))
a = val_extr(sheet_tod.cell(40, 13))
document.add_paragraph('- исполнительное производство – {} млн. руб. (переданы в ССП и Банк на исполнение).'.format(a))
a = val_extr(sheet_tod.cell(40, 14))
b = val_extr(sheet_tod.cell(40, 15))
c = val_extr(sheet_tod.cell(40, 16))
document.add_paragraph('Безнадежная задолженность составляет {0} млн. руб. ({1} % от суммы задолженности), в том числе, вынесены акты о невозможности взыскания на сумму {2} млн. руб.'.format(a, b, c))
document.add_paragraph(''.format())

#сохранение документа
itog_spr = 'Справка 100 тыс. ' + date_tod + '.docx'
document.save(itog_spr)

#сначала строка потом столбец
#print(val_extr(sheet_tod.cell(29, 3)))